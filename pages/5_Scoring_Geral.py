# pages/5_Scoring_Geral.py — Tabela de Scoring Geral

import streamlit as st
import json
import os
import sys
import io
from datetime import datetime

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config
from core.styles import inject_css, page_header, skill_tags, score_badge, footer, LOGO_SVG
from qa_agent import QAAgent

st.set_page_config(page_title="Scoring Geral &middot; TalentIQ", page_icon="&#128200;", layout="wide")
inject_css()

# â”€â”€ Data helpers â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _load_data():
    path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data", "jobs.json")
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {"vagas": [], "candidatos": []}


def _build_df(candidatos: list, vagas: list) -> pd.DataFrame:
    vaga_map = {v["id"]: v for v in vagas}
    rows = []
    for c in candidatos:
        fit = c.get("fit_resultado", {})
        det = fit.get("pontuacao_detalhada", {})
        vaga = vaga_map.get(c["vaga_id"], {})
        rows.append({
            "ID": c["id"],
            "Candidato": c["nome"],
            "Email": c.get("email", "—"),
            "Vaga": c["vaga_titulo"],
            "Departamento": vaga.get("departamento", "—"),
            "Etapa": c.get("etapa", "—"),
            "Score Total": c.get("score_fit"),
            "Competências (50)": det.get("competencias"),
            "Experiência (30)": det.get("experiencia"),
            "Formação (20)": det.get("formacao"),
            "Nível Alinhamento": fit.get("nivel_alinhamento", "Não calculado"),
            "Motor": fit.get("metodo", "—"),
            "Data Candidatura": c.get("data_candidatura", "—"),
            "Exp. Anos": c["perfil"].get("experiencia_anos", 0),
            "NÂº Competências": len(c["perfil"].get("competencias", [])),
        })
    return pd.DataFrame(rows)


# â”€â”€ Excel export â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _export_excel(df: pd.DataFrame, vagas: list) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import (PatternFill, Font, Alignment,
                                  Border, Side, GradientFill)
    from openpyxl.utils import get_column_letter
    from openpyxl.chart import BarChart, Reference

    wb = Workbook()

    # â”€â”€ Sheet 1: Ranking Geral â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws1 = wb.active
    ws1.title = "Ranking Geral"

    # Header block
    ws1.merge_cells("A1:O1")
    ws1["A1"] = "TalentIQ — Tabela de Scoring Geral"
    ws1["A1"].font = Font(bold=True, size=16, color="FFFFFF")
    ws1["A1"].fill = PatternFill("solid", fgColor="1D4ED8")
    ws1["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws1.row_dimensions[1].height = 32

    ws1.merge_cells("A2:O2")
    ws1["A2"] = f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}  &middot;  Motor IA: {config.LLM_ENGINE.upper()}"
    ws1["A2"].font = Font(italic=True, size=10, color="64748B")
    ws1["A2"].alignment = Alignment(horizontal="center")

    # Column headers
    cols = [
        "Rank", "Candidato", "Email", "Vaga", "Departamento", "Etapa",
        "Score Total", "Competências\n(/50)", "Experiência\n(/30)", "Formação\n(/20)",
        "Nível Alinhamento", "Motor IA", "Data Candidatura", "Exp. Anos", "NÂº Competências"
    ]
    col_widths = [6, 24, 28, 28, 18, 20, 12, 14, 14, 12, 22, 20, 18, 10, 16]

    header_fill = PatternFill("solid", fgColor="1E3A8A")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    thin = Side(style="thin", color="CBD5E1")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col_idx, (col_name, width) in enumerate(zip(cols, col_widths), 1):
        cell = ws1.cell(row=3, column=col_idx, value=col_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
        ws1.column_dimensions[get_column_letter(col_idx)].width = width
    ws1.row_dimensions[3].height = 36

    # Sort by score
    df_sorted = df.sort_values("Score Total", ascending=False, na_position="last").reset_index(drop=True)

    # Score colour fills
    fill_high   = PatternFill("solid", fgColor="DCFCE7")  # green
    fill_mid    = PatternFill("solid", fgColor="FEF9C3")  # yellow
    fill_low    = PatternFill("solid", fgColor="FEE2E2")  # red
    fill_none   = PatternFill("solid", fgColor="F8FAFC")
    font_normal = Font(size=10)
    font_bold   = Font(size=10, bold=True)

    for row_idx, row in df_sorted.iterrows():
        excel_row = row_idx + 4
        score = row["Score Total"]
        if score is None:
            row_fill = fill_none
        elif score >= 75:
            row_fill = fill_high
        elif score >= 50:
            row_fill = fill_mid
        else:
            row_fill = fill_low

        values = [
            row_idx + 1,
            row["Candidato"],
            row["Email"],
            row["Vaga"],
            row["Departamento"],
            row["Etapa"],
            score,
            row["Competências (50)"],
            row["Experiência (30)"],
            row["Formação (20)"],
            row["Nível Alinhamento"],
            row["Motor"],
            row["Data Candidatura"],
            row["Exp. Anos"],
            row["NÂº Competências"],
        ]
        for col_idx, val in enumerate(values, 1):
            cell = ws1.cell(row=excel_row, column=col_idx, value=val)
            cell.fill = row_fill
            cell.font = font_bold if col_idx in (1, 7) else font_normal
            cell.alignment = Alignment(horizontal="center" if col_idx in (1, 7, 8, 9, 10, 14, 15) else "left",
                                       vertical="center")
            cell.border = border
        ws1.row_dimensions[excel_row].height = 18

    ws1.freeze_panes = "A4"
    ws1.auto_filter.ref = f"A3:O{3 + len(df_sorted)}"

    # â”€â”€ Sheet 2: Por Vaga â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws2 = wb.create_sheet("Por Vaga")
    ws2.merge_cells("A1:F1")
    ws2["A1"] = "Resumo por Vaga"
    ws2["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws2["A1"].fill = PatternFill("solid", fgColor="1D4ED8")
    ws2["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws2.row_dimensions[1].height = 28

    vaga_headers = ["Vaga", "Candidatos", "Score Médio", "Score Máx", "Score Mín", "Nível Dominante"]
    vaga_widths  = [32, 12, 14, 12, 12, 24]
    for col_idx, (h, w) in enumerate(zip(vaga_headers, vaga_widths), 1):
        cell = ws2.cell(row=2, column=col_idx, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border
        ws2.column_dimensions[get_column_letter(col_idx)].width = w

    vaga_groups = df.groupby("Vaga")
    for r_idx, (vaga_nome, grp) in enumerate(vaga_groups, 3):
        scores_v = grp["Score Total"].dropna()
        nivel_counts = grp["Nível Alinhamento"].value_counts()
        nivel_dom = nivel_counts.index[0] if len(nivel_counts) else "—"
        row_data = [
            vaga_nome,
            len(grp),
            round(scores_v.mean(), 1) if len(scores_v) else "—",
            int(scores_v.max()) if len(scores_v) else "—",
            int(scores_v.min()) if len(scores_v) else "—",
            nivel_dom,
        ]
        for c_idx, val in enumerate(row_data, 1):
            cell = ws2.cell(row=r_idx, column=c_idx, value=val)
            cell.font = font_normal
            cell.alignment = Alignment(horizontal="center" if c_idx > 1 else "left", vertical="center")
            cell.border = border
        ws2.row_dimensions[r_idx].height = 18

    # â”€â”€ Sheet 3: Matriz Competências â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    ws3 = wb.create_sheet("Competências")
    ws3.merge_cells("A1:B1")
    ws3["A1"] = "Competências mais frequentes nos candidatos"
    ws3["A1"].font = Font(bold=True, size=12, color="FFFFFF")
    ws3["A1"].fill = PatternFill("solid", fgColor="1D4ED8")
    ws3["A1"].alignment = Alignment(horizontal="center")
    ws3.row_dimensions[1].height = 24

    all_comps: dict = {}
    for c in _load_data().get("candidatos", []):
        for comp in c["perfil"].get("competencias", []):
            all_comps[comp] = all_comps.get(comp, 0) + 1

    ws3.cell(row=2, column=1, value="Competência").font = header_font
    ws3.cell(row=2, column=1).fill = header_fill
    ws3.cell(row=2, column=2, value="Frequência").font = header_font
    ws3.cell(row=2, column=2).fill = header_fill
    ws3.column_dimensions["A"].width = 30
    ws3.column_dimensions["B"].width = 14

    for r_idx, (comp, freq) in enumerate(sorted(all_comps.items(), key=lambda x: -x[1]), 3):
        ws3.cell(row=r_idx, column=1, value=comp).font = font_normal
        ws3.cell(row=r_idx, column=2, value=freq).font = font_normal
        ws3.cell(row=r_idx, column=2).alignment = Alignment(horizontal="center")

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


# â”€â”€ Word report â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def _export_word(df: pd.DataFrame, vagas: list, candidatos: list) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TalentIQ")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(29, 78, 216)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Relatório de Scoring Geral de Candidatos").font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta.add_run(
        f"Gerado em: {datetime.now().strftime('%d de %B de %Y, %H:%M')}  &middot;  "
        f"Motor IA: {config.LLM_ENGINE.upper()}"
    )
    run_meta.font.size = Pt(10)
    run_meta.font.color.rgb = RGBColor(100, 116, 139)
    run_meta.font.italic = True

    doc.add_paragraph()

    # â”€â”€ Summary stats â”€â”€
    p = doc.add_heading("1. Resumo Executivo", level=1)
    p.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    df_scored = df[df["Score Total"].notna()]
    n_total   = len(df)
    n_scored  = len(df_scored)
    n_altos   = len(df_scored[df_scored["Score Total"] >= 75])
    n_medios  = len(df_scored[(df_scored["Score Total"] >= 50) & (df_scored["Score Total"] < 75)])
    n_baixos  = len(df_scored[df_scored["Score Total"] < 50])
    avg_score = round(df_scored["Score Total"].mean(), 1) if n_scored else "—"

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_s = ["Total Candidatos", "Com Score", "Alto (â‰¥75)", "Médio (50-74)", "Score Médio"]
    values_s  = [str(n_total), str(n_scored), str(n_altos), str(n_medios), str(avg_score)]
    for i, (h, v) in enumerate(zip(headers_s, values_s)):
        hcell = tbl.rows[0].cells[i]
        vcell = tbl.rows[1].cells[i]
        hcell.text = h
        hcell.paragraphs[0].runs[0].font.bold = True
        hcell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_bg(hcell, "1E3A8A")
        hcell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        vcell.text = v
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        vcell.paragraphs[0].runs[0].font.bold = True
        vcell.paragraphs[0].runs[0].font.size = Pt(12)
        set_cell_bg(vcell, "EFF6FF")

    doc.add_paragraph()

    # â”€â”€ Ranking table â”€â”€
    p2 = doc.add_heading("2. Ranking Geral de Candidatos", level=1)
    p2.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    df_sorted = df.sort_values("Score Total", ascending=False, na_position="last").reset_index(drop=True)

    rank_cols = ["Rank", "Candidato", "Vaga", "Score", "Competências", "Experiência", "Formação", "Nível"]
    tbl2 = doc.add_table(rows=1 + len(df_sorted), cols=len(rank_cols))
    tbl2.style = "Table Grid"

    for i, h in enumerate(rank_cols):
        cell = tbl2.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_bg(cell, "1E3A8A")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for r_idx, row in df_sorted.iterrows():
        score = row["Score Total"]
        bg = ("DCFCE7" if score and score >= 75 else
              "FEF9C3" if score and score >= 50 else
              "FEE2E2" if score else "F8FAFC")
        vals = [
            str(r_idx + 1),
            row["Candidato"],
            row["Vaga"][:30] + ("â€¦" if len(row["Vaga"]) > 30 else ""),
            str(int(score)) if score is not None else "—",
            str(row["Competências (50)"]) + "/50" if row["Competências (50)"] is not None else "—",
            str(row["Experiência (30)"]) + "/30" if row["Experiência (30)"] is not None else "—",
            str(row["Formação (20)"]) + "/20" if row["Formação (20)"] is not None else "—",
            row["Nível Alinhamento"],
        ]
        for c_idx, val in enumerate(vals):
            cell = tbl2.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.CENTER
                                            if c_idx in (0, 3, 4, 5, 6) else WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_bg(cell, bg)

    doc.add_paragraph()

    # â”€â”€ Per-job sections â”€â”€
    p3 = doc.add_heading("3. Análise Detalhada por Vaga", level=1)
    p3.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    vaga_map = {v["id"]: v for v in vagas}

    for vaga_nome, grp in df.groupby("Vaga"):
        doc.add_heading(f"&#128204; {vaga_nome}", level=2)
        grp_s = grp.sort_values("Score Total", ascending=False, na_position="last")

        scores_v = grp_s["Score Total"].dropna()
        if len(scores_v):
            stats_p = doc.add_paragraph()
            stats_p.add_run(
                f"Candidatos: {len(grp_s)}  &middot;  Score médio: {round(scores_v.mean(),1)}  &middot;  "
                f"Máximo: {int(scores_v.max())}  &middot;  Mínimo: {int(scores_v.min())}"
            ).font.size = Pt(10)

        for _, row in grp_s.iterrows():
            score = row["Score Total"]
            nivel = row["Nível Alinhamento"]
            icon  = "ðŸŸ¢" if score and score >= 75 else "ðŸŸ¡" if score and score >= 50 else "ðŸ”´"
            p_cand = doc.add_paragraph(style="List Bullet")
            run_n = p_cand.add_run(f"{row['Candidato']}")
            run_n.bold = True
            run_n.font.size = Pt(10)
            run_s = p_cand.add_run(
                f"  —  Score: {int(score) if score is not None else '—'}/100  &middot;  {nivel}  &middot;  Etapa: {row['Etapa']}"
            )
            run_s.font.size = Pt(10)
            run_s.font.color.rgb = RGBColor(100, 116, 139)

        doc.add_paragraph()

    # â”€â”€ Recommendations â”€â”€
    p4 = doc.add_heading("4. Recomendações", level=1)
    p4.runs[0].font.color.rgb = RGBColor(30, 58, 138)

    top3 = df_sorted[df_sorted["Score Total"].notna()].head(3)
    for _, row in top3.iterrows():
        p_rec = doc.add_paragraph(style="List Bullet")
        run_r = p_rec.add_run(
            f"{row['Candidato']} ({row['Vaga']}) — Score {int(row['Score Total'])}/100: "
            f"candidato de {row['Nível Alinhamento'].lower()} para progressão imediata."
        )
        run_r.font.size = Pt(10)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_p.add_run(
        f"TalentIQ v{config.APP_VERSION}  &middot;  Documento gerado automaticamente  &middot;  "
        f"{datetime.now().strftime('%d/%m/%Y')}"
    )
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = RGBColor(148, 163, 184)
    run_f.font.italic = True

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# PAGE RENDER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

qa = QAAgent(config.APP_NAME, config.APP_VERSION)
qa.display_qa_dashboard(qa.run_full_qa_suite())

_sb_logo = LOGO_SVG.replace("\n", "").replace("  ", " ")
with st.sidebar:
    st.markdown(
        '<div style="text-align:center;padding:4px 0 12px;">'
        '<div style="display:inline-block;filter:drop-shadow(0 0 10px rgba(29,78,216,0.5));">'
        + _sb_logo +
        '</div>'
        '<div style="font-size:1.5rem;font-weight:800;color:#ffffff;margin-top:6px;">'
        'Talent<span style="color:#3b82f6;">IQ</span></div>'
        '</div>',
        unsafe_allow_html=True,
    )

page_header("&#128200;", "Scoring Geral",
            "Vista consolidada de todos os candidatos &middot; Exportação Excel &middot; Relatório Word")

dados = _load_data()
vagas      = dados.get("vagas", [])
candidatos = dados.get("candidatos", [])

if not candidatos:
    st.markdown("""
    <div style="background:#eff6ff; border:1px solid #bfdbfe; border-radius:12px;
                padding:24px; text-align:center; color:#1d4ed8;">
        &#8505; Nenhum candidato registado ainda. Carregue CVs em <strong>&#128196; Carregar CV</strong> primeiro.
    </div>
    """, unsafe_allow_html=True)
    st.stop()

df_all = _build_df(candidatos, vagas)
df_scored = df_all[df_all["Score Total"].notna()]

# â”€â”€ KPI row â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
scores_all = df_scored["Score Total"].tolist()
n_altos  = len([s for s in scores_all if s >= 75])
n_medios = len([s for s in scores_all if 50 <= s < 75])
n_baixos = len([s for s in scores_all if s < 50])

col1, col2, col3, col4, col5, col6 = st.columns(6)
col1.metric("&#128101; Candidatos", len(df_all))
col2.metric("&#128203; Vagas", len(vagas))
col3.metric("&#127919; Com Score", len(df_scored))
col4.metric("ðŸŸ¢ Alto Fit", n_altos)
col5.metric("ðŸŸ¡ Médio Fit", n_medios)
col6.metric("&#128202; Score Médio", f"{round(sum(scores_all)/len(scores_all))}%" if scores_all else "—")

st.markdown("<br>", unsafe_allow_html=True)

# â”€â”€ Tabs â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
tab1, tab2, tab3, tab4 = st.tabs([
    "&#128202;  Vista Consolidada",
    "&#128200;  Gráficos",
    "&#128229;  Exportar Excel",
    "&#128196;  Relatório Word",
])

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TAB 1 — Vista consolidada
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
with tab1:
    st.markdown("<br>", unsafe_allow_html=True)

    # Filters
    col_f1, col_f2, col_f3 = st.columns(3)
    with col_f1:
        vaga_filter = st.multiselect(
            "&#128269; Filtrar por Vaga",
            options=df_all["Vaga"].unique().tolist(),
            default=[],
            placeholder="Todas as vagas"
        )
    with col_f2:
        nivel_filter = st.multiselect(
            "&#127919; Nível de Alinhamento",
            options=["Alto Alinhamento", "Alinhamento Médio", "Baixo Alinhamento", "Não calculado"],
            default=[],
            placeholder="Todos os níveis"
        )
    with col_f3:
        etapa_filter = st.multiselect(
            "ðŸ“ Etapa no Pipeline",
            options=df_all["Etapa"].unique().tolist(),
            default=[],
            placeholder="Todas as etapas"
        )

    df_view = df_all.copy()
    if vaga_filter:
        df_view = df_view[df_view["Vaga"].isin(vaga_filter)]
    if nivel_filter:
        df_view = df_view[df_view["Nível Alinhamento"].isin(nivel_filter)]
    if etapa_filter:
        df_view = df_view[df_view["Etapa"].isin(etapa_filter)]

    df_view_sorted = df_view.sort_values("Score Total", ascending=False, na_position="last").reset_index(drop=True)
    df_view_sorted.insert(0, "Rank", df_view_sorted.index + 1)

    st.markdown(f"**{len(df_view_sorted)} candidato(s)** encontrado(s)")

    # Render coloured table
    def colour_score(val):
        if val is None or val == "":
            return "color: #94a3b8"
        try:
            v = float(val)
            if v >= 75:
                return "background-color: #dcfce7; color: #15803d; font-weight:700"
            if v >= 50:
                return "background-color: #fef9c3; color: #a16207; font-weight:700"
            return "background-color: #fee2e2; color: #b91c1c; font-weight:700"
        except Exception:
            return ""

    display_cols = ["Rank", "Candidato", "Vaga", "Etapa",
                    "Score Total", "Competências (50)", "Experiência (30)", "Formação (20)",
                    "Nível Alinhamento", "Exp. Anos"]

    styled = (
        df_view_sorted[display_cols]
        .style
        .map(colour_score, subset=["Score Total"])
        .format({"Score Total": lambda x: f"{int(x)}/100" if x is not None else "—",
                 "Competências (50)": lambda x: f"{int(x)}/50" if x is not None else "—",
                 "Experiência (30)": lambda x: f"{int(x)}/30" if x is not None else "—",
                 "Formação (20)": lambda x: f"{int(x)}/20" if x is not None else "—"}, na_rep="—")
        .set_properties(**{"font-size": "13px"})
    )

    st.dataframe(styled, use_container_width=True, hide_index=True, height=500)

    # Candidate detail expander
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("#### ðŸ”Ž Detalhe do Candidato")
    nomes_map = {f"{r['Candidato']} &middot; {r['Vaga']} (Score: {r['Score Total']})": i
                 for i, r in df_view_sorted.iterrows()}
    if nomes_map:
        sel_key = st.selectbox("Seleccionar candidato", list(nomes_map.keys()))
        sel_row = df_view_sorted.iloc[nomes_map[sel_key]]
        cand_obj = next((c for c in candidatos if c["id"] == sel_row["ID"]), None)

        if cand_obj:
            fit = cand_obj.get("fit_resultado", {})
            col_a, col_b, col_c, col_d = st.columns(4)
            score = cand_obj.get("score_fit")
            col_a.metric("&#127919; Score", f"{score}/100" if score is not None else "—")
            col_b.metric("&#129504; Competências", f"{fit.get('pontuacao_detalhada',{}).get('competencias','—')}/50")
            col_c.metric("&#9203; Experiência", f"{fit.get('pontuacao_detalhada',{}).get('experiencia','—')}/30")
            col_d.metric("&#127891; Formação", f"{fit.get('pontuacao_detalhada',{}).get('formacao','—')}/20")

            if fit.get("explicacao"):
                st.markdown("**&#128203; Análise:**")
                for linha in fit["explicacao"]:
                    st.markdown(f"<div style='padding:3px 0; color:#374151;'>{linha}</div>",
                                unsafe_allow_html=True)

            col_e, col_f = st.columns(2)
            with col_e:
                st.markdown("**&#129504; Competências do Candidato:**")
                comps = cand_obj["perfil"].get("competencias", [])
                vaga_obj = next((v for v in vagas if v["id"] == cand_obj["vaga_id"]), {})
                comps_vaga = vaga_obj.get("competencias_requeridas", [])
                matches = [c for c in comps if c.lower() in [v.lower() for v in comps_vaga]]
                skill_tags(comps, matches=matches)
            with col_f:
                st.markdown("**&#10060; Lacunas:**")
                gaps = [c for c in comps_vaga if c.lower() not in [x.lower() for x in comps]]
                if gaps:
                    skill_tags(gaps, gaps=gaps)
                else:
                    st.markdown('<span class="skill-tag-match">&#10003; Sem lacunas</span>',
                                unsafe_allow_html=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TAB 2 — Charts
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
with tab2:
    st.markdown("<br>", unsafe_allow_html=True)

    if df_scored.empty:
        st.info("Calcule os scores em **&#127919; Pontuação Fit** para ver os gráficos.")
    else:
        col_g1, col_g2 = st.columns(2)

        with col_g1:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown("#### &#127942; Top 10 Candidatos")
            top10 = df_scored.nlargest(10, "Score Total")
            fig_bar = px.bar(
                top10, x="Score Total", y="Candidato",
                orientation="h", color="Score Total",
                color_continuous_scale=["#ef4444", "#f59e0b", "#10b981"],
                range_color=[0, 100],
                text="Score Total",
            )
            fig_bar.update_traces(texttemplate="%{text}/100", textposition="outside")
            fig_bar.update_layout(
                yaxis={"categoryorder": "total ascending"},
                showlegend=False, coloraxis_showscale=False,
                margin=dict(l=10, r=40, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Inter", color="#1e3a8a"),
            )
            st.plotly_chart(fig_bar, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        with col_g2:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown("#### &#127919; Distribuição de Scores")
            fig_hist = px.histogram(
                df_scored, x="Score Total", nbins=10,
                color_discrete_sequence=["#1d4ed8"],
                labels={"Score Total": "Score", "count": "Candidatos"},
            )
            fig_hist.update_layout(
                margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Inter", color="#1e3a8a"),
            )
            st.plotly_chart(fig_hist, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        col_g3, col_g4 = st.columns(2)

        with col_g3:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown("#### &#128202; Score Médio por Vaga")
            score_por_vaga = (
                df_scored.groupby("Vaga")["Score Total"].mean().round(1).reset_index()
                .sort_values("Score Total", ascending=False)
            )
            fig_vaga = px.bar(
                score_por_vaga, x="Vaga", y="Score Total",
                color="Score Total",
                color_continuous_scale=["#ef4444", "#f59e0b", "#10b981"],
                range_color=[0, 100], text="Score Total",
            )
            fig_vaga.update_traces(texttemplate="%{text:.1f}", textposition="outside")
            fig_vaga.update_layout(
                showlegend=False, coloraxis_showscale=False,
                margin=dict(l=10, r=10, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Inter", color="#1e3a8a"),
                xaxis_tickangle=-20,
            )
            st.plotly_chart(fig_vaga, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

        with col_g4:
            st.markdown('<div class="section-card">', unsafe_allow_html=True)
            st.markdown("#### ðŸ•¸ Radar — Dimensões do Score")
            top5 = df_scored.nlargest(5, "Score Total")
            categories = ["Competências (50)", "Experiência (30)", "Formação (20)"]
            fig_radar = go.Figure()
            for _, row in top5.iterrows():
                vals = [
                    (row["Competências (50)"] or 0) / 50 * 100,
                    (row["Experiência (30)"] or 0) / 30 * 100,
                    (row["Formação (20)"] or 0) / 20 * 100,
                ]
                vals_closed = vals + [vals[0]]
                fig_radar.add_trace(go.Scatterpolar(
                    r=vals_closed,
                    theta=["Competências", "Experiência", "Formação", "Competências"],
                    fill="toself", name=row["Candidato"][:18],
                    opacity=0.6,
                ))
            fig_radar.update_layout(
                polar=dict(radialaxis=dict(visible=True, range=[0, 100])),
                margin=dict(l=30, r=30, t=30, b=30),
                paper_bgcolor="rgba(0,0,0,0)",
                font=dict(family="Inter", color="#1e3a8a"),
                legend=dict(font=dict(size=9)),
            )
            st.plotly_chart(fig_radar, use_container_width=True)
            st.markdown('</div>', unsafe_allow_html=True)

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TAB 3 — Export Excel
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
with tab3:
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("#### &#128229; Exportar para Excel")
    st.markdown("""
    O ficheiro Excel inclui **3 folhas**:
    - **Ranking Geral** — todos os candidatos ordenados por score, com código de cores automático
    - **Por Vaga** — resumo estatístico por vaga (média, máximo, mínimo)
    - **Competências** — frequência de competências identificadas nos candidatos
    """)

    col_ex1, col_ex2 = st.columns([1, 2])
    with col_ex1:
        nome_ficheiro = st.text_input(
            "Nome do ficheiro",
            value=f"TalentIQ_Scoring_{datetime.now().strftime('%Y%m%d')}",
        )

    if st.button("&#128229;  Gerar e Descarregar Excel", type="primary", use_container_width=False):
        with st.spinner("A gerar ficheiro Excel..."):
            excel_bytes = _export_excel(df_all, vagas)
        st.download_button(
            label="â¬‡  Descarregar Excel",
            data=excel_bytes,
            file_name=f"{nome_ficheiro}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        st.success("&#9989; Ficheiro pronto para descarregar!")

    st.markdown('</div>', unsafe_allow_html=True)

    # Preview table
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("**Pré-visualização dos dados a exportar:**")
    st.dataframe(
        df_all[["Candidato", "Vaga", "Score Total", "Competências (50)",
                "Experiência (30)", "Formação (20)", "Nível Alinhamento", "Etapa"]],
        use_container_width=True,
        hide_index=True,
    )

# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TAB 4 — Word Report
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
with tab4:
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown('<div class="section-card">', unsafe_allow_html=True)
    st.markdown("#### &#128196; Relatório Word")
    st.markdown("""
    O relatório Word inclui:
    - **Resumo Executivo** — KPIs gerais de scoring
    - **Ranking Geral** — tabela completa de candidatos com código de cores
    - **Análise por Vaga** — detalhe de candidatos por posição
    - **Recomendações** — top 3 candidatos com justificação automática
    """)

    col_w1, col_w2 = st.columns([1, 2])
    with col_w1:
        nome_relatorio = st.text_input(
            "Nome do relatório",
            value=f"TalentIQ_Relatorio_{datetime.now().strftime('%Y%m%d')}",
        )

    if st.button("&#128196;  Gerar e Descarregar Relatório Word", type="primary"):
        with st.spinner("A gerar relatório Word..."):
            word_bytes = _export_word(df_all, vagas, candidatos)
        st.download_button(
            label="â¬‡  Descarregar Relatório (.docx)",
            data=word_bytes,
            file_name=f"{nome_relatorio}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.success("&#9989; Relatório pronto para descarregar!")

    st.markdown('</div>', unsafe_allow_html=True)

footer(config.APP_VERSION, config.LLM_ENGINE)


