# gerar_documentos_teste.py
# Generates mock CVs and a ToR for TalentIQ testing
# Run: python gerar_documentos_teste.py

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), "documentos_teste")
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p


def add_section(doc, title, content_lines):
    set_heading(doc, title, level=2, color=(31, 73, 125))
    for line in content_lines:
        doc.add_paragraph(line, style="List Bullet" if line.startswith("•") else "Normal")


# ─────────────────────────────────────────────
# CV 1 — Strong fit: Data Analyst (Senior)
# ─────────────────────────────────────────────
def criar_cv_ana_macie():
    doc = Document()
    doc.add_heading("ANA BEATRIZ MACIE", 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("📧 ana.macie@gmail.com  |  📞 +258 84 123 4567  |  📍 Maputo, Moçambique")

    doc.add_paragraph()

    add_section(doc, "Resumo Profissional", [
        "Analista de Dados Sénior com 7 anos de experiência em análise estatística, "
        "visualização de dados e suporte à tomada de decisão em contextos de saúde pública. "
        "Domínio avançado de Python, SQL e Power BI. Experiência comprovada em projectos "
        "financiados por USAID, PEPFAR e Jhpiego em Moçambique e Angola."
    ])

    add_section(doc, "Competências Técnicas", [
        "• Python (pandas, numpy, scikit-learn, matplotlib)",
        "• SQL — PostgreSQL, MySQL",
        "• Power BI, Tableau, Excel Avançado",
        "• Machine Learning, Data Science",
        "• DHIS2, SISMA, OpenMRS",
        "• Gestão de Projetos, Scrum, Agile",
        "• R (estatística descritiva e inferencial)",
    ])

    add_section(doc, "Experiência Profissional", [
        "Analista de Dados Sénior | Jhpiego Moçambique | 2019 – 2026",
        "• Desenvolvimento de dashboards de monitoria para programas de HIV/SIDA",
        "• Análise de dados de 450+ unidades sanitárias nas províncias de Zambézia e Nampula",
        "• Formação de equipas locais em uso de DHIS2 e Power BI",
        "",
        "Analista de Dados | Ministério da Saúde — MISAU | 2017 – 2019",
        "• Processamento e validação de dados do Boletim Epidemiológico Semanal",
        "• Suporte técnico à implementação do SISMA a nível provincial",
    ])

    add_section(doc, "Formação Académica", [
        "Mestrado em Saúde Pública — Epidemiologia | Universidade Eduardo Mondlane | 2017",
        "Licenciatura em Estatística | Universidade Eduardo Mondlane | 2015",
    ])

    add_section(doc, "Idiomas", [
        "• Português — Nativo",
        "• Inglês — Fluente (C1)",
        "• Changana — Nativo",
    ])

    path = os.path.join(OUTPUT_DIR, "CV_Ana_Macie_Analista_Senior.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# CV 2 — Medium fit: IT Manager
# ─────────────────────────────────────────────
def criar_cv_carlos_sitoe():
    doc = Document()
    doc.add_heading("CARLOS MANUEL SITOE", 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("📧 c.sitoe@outlook.com  |  📞 +258 82 987 6543  |  📍 Beira, Moçambique")

    doc.add_paragraph()

    add_section(doc, "Resumo Profissional", [
        "Gestor de Tecnologias de Informação com 5 anos de experiência em administração "
        "de sistemas, suporte técnico e gestão de redes. Conhecimento moderado de SQL e "
        "Excel. Interesse crescente em análise de dados e visualização. Experiência em "
        "gestão de equipas pequenas."
    ])

    add_section(doc, "Competências Técnicas", [
        "• SQL (nível intermédio)",
        "• Excel (intermédio)",
        "• Redes e Infraestrutura de TI",
        "• Suporte técnico e helpdesk",
        "• Gestão de equipa",
        "• Windows Server, Active Directory",
    ])

    add_section(doc, "Experiência Profissional", [
        "Gestor de TI | Banco Comercial e de Investimentos (BCI) | 2021 – 2026",
        "• Gestão da infraestrutura de TI de 3 balcões na cidade da Beira",
        "• Coordenação de equipa de 4 técnicos de suporte",
        "• Implementação de sistema de backup e recuperação de dados",
        "",
        "Técnico de Suporte de TI | Vodacom Moçambique | 2019 – 2021",
        "• Resolução de incidentes técnicos de clientes corporativos",
        "• Configuração de redes LAN/WAN em pequenas empresas",
    ])

    add_section(doc, "Formação Académica", [
        "Licenciatura em Engenharia Informática | Instituto Superior de Ciências e Tecnologia | 2019",
        "Certificação CCNA — Cisco Networking | 2020",
    ])

    add_section(doc, "Idiomas", [
        "• Português — Nativo",
        "• Inglês — Intermediário (B1)",
        "• Sena — Nativo",
    ])

    path = os.path.join(OUTPUT_DIR, "CV_Carlos_Sitoe_Gestor_TI.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# CV 3 — Weak fit: Teacher (career change)
# ─────────────────────────────────────────────
def criar_cv_fatima_cumbe():
    doc = Document()
    doc.add_heading("FÁTIMA LOURENÇO CUMBE", 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("📧 fatima.cumbe@yahoo.com  |  📞 +258 86 555 1234  |  📍 Nampula, Moçambique")

    doc.add_paragraph()

    add_section(doc, "Resumo Profissional", [
        "Professora de Matemática e Ciências com 8 anos de experiência no ensino secundário. "
        "Forte capacidade de comunicação e liderança de turmas. Interesse em transição de "
        "carreira para análise de dados. Conhecimentos básicos de Excel adquiridos em "
        "auto-formação nos últimos 6 meses."
    ])

    add_section(doc, "Competências", [
        "• Ensino e formação",
        "• Comunicação e apresentação",
        "• Excel (nível básico — auto-didacta)",
        "• Liderança de grupos",
        "• Matemática e Estatística",
        "• Microsoft Office (Word, PowerPoint)",
    ])

    add_section(doc, "Experiência Profissional", [
        "Professora de Matemática | Escola Secundária de Nampula | 2018 – 2026",
        "• Leccionação de Matemática, Física e Estatística ao nível do 10.º ao 12.º ano",
        "• Coordenação do clube de Ciências e Tecnologia da escola",
        "• Preparação de alunos para exames nacionais — taxa de aprovação de 82%",
        "",
        "Explicadora Particular | Independente | 2016 – 2018",
        "• Aulas particulares de Matemática para alunos do ensino básico e secundário",
    ])

    add_section(doc, "Formação Académica", [
        "Licenciatura em Ensino de Matemática | Universidade Pedagógica de Nampula | 2016",
        "Curso Técnico em Informática Básica | Instituto Médio de Nampula | 2012",
    ])

    add_section(doc, "Idiomas", [
        "• Português — Nativo",
        "• Makua — Nativo",
        "• Inglês — Básico (A2)",
    ])

    path = os.path.join(OUTPUT_DIR, "CV_Fatima_Cumbe_Professora.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# CV 4 — Strong fit: HR Recruiter
# ─────────────────────────────────────────────
def criar_cv_joao_nhantumbo():
    doc = Document()
    doc.add_heading("JOÃO PEDRO NHANTUMBO", 0)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("📧 jp.nhantumbo@gmail.com  |  📞 +258 84 777 8899  |  📍 Maputo, Moçambique")

    doc.add_paragraph()

    add_section(doc, "Resumo Profissional", [
        "Especialista em Recursos Humanos com 6 anos de experiência em recrutamento e "
        "seleção, gestão de desempenho e desenvolvimento organizacional. Sólida experiência "
        "em entrevistas por competências, gestão de pipeline de candidatos e onboarding. "
        "Trabalhou em organizações nacionais e internacionais no sector de saúde."
    ])

    add_section(doc, "Competências", [
        "• Recrutamento e seleção",
        "• Gestão de equipa",
        "• Entrevistas por competências",
        "• Gestão de desempenho",
        "• Recursos humanos",
        "• Excel Avançado, Power BI",
        "• Negociação e comunicação",
        "• Formação e treinamento",
        "• Liderança",
    ])

    add_section(doc, "Experiência Profissional", [
        "Especialista de RH Sénior | FHI 360 Moçambique | 2021 – 2026",
        "• Liderança do processo de recrutamento para 60+ posições por ano",
        "• Desenvolvimento de políticas de RH alinhadas com legislação laboral moçambicana",
        "• Gestão do ciclo completo de avaliação de desempenho para 120 colaboradores",
        "• Coordenação de programas de formação e desenvolvimento de capacidades",
        "",
        "Técnico de RH | Chemonics International | 2018 – 2021",
        "• Processamento de folha de salários e gestão de benefícios",
        "• Apoio na organização de processos de selecção e entrevistas",
        "• Gestão de arquivos e documentação de colaboradores",
        "",
        "Assistente Administrativo | Ministério da Saúde | 2016 – 2018",
        "• Apoio administrativo à Direcção de Recursos Humanos",
    ])

    add_section(doc, "Formação Académica", [
        "MBA em Gestão de Recursos Humanos | ISEG — Instituto Superior de Economia e Gestão | 2021",
        "Licenciatura em Gestão de Empresas | Universidade Politécnica de Maputo | 2016",
    ])

    add_section(doc, "Idiomas", [
        "• Português — Nativo",
        "• Inglês — Fluente (C1)",
        "• Espanhol — Intermediário (B2)",
        "• Ronga — Nativo",
    ])

    path = os.path.join(OUTPUT_DIR, "CV_Joao_Nhantumbo_RH_Senior.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# ToR — Analista de Dados Sénior
# ─────────────────────────────────────────────
def criar_tor_analista_dados():
    doc = Document()

    # Title
    title = doc.add_heading("TERMOS DE REFERÊNCIA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading("Analista de Dados Sénior — Programa de Saúde Pública", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Org info
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    cells = [
        ("Organização", "TalentIQ Health Partners"),
        ("Localização", "Maputo, Moçambique (com viagens provinciais)"),
        ("Duração do Contrato", "12 meses renováveis"),
        ("Data Limite de Candidatura", "31 de Julho de 2026"),
    ]
    for i, (label, value) in enumerate(cells):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Background
    set_heading(doc, "1. Contexto", level=2, color=(31, 73, 125))
    doc.add_paragraph(
        "A TalentIQ Health Partners é uma organização de implementação que apoia o "
        "Ministério da Saúde de Moçambique na monitoria e avaliação de programas de "
        "HIV/SIDA, Saúde Materna e Saúde Infantil financiados pelo PEPFAR e USAID. "
        "A organização opera em 5 províncias com cobertura de mais de 600 unidades sanitárias."
    )
    doc.add_paragraph(
        "No âmbito da expansão das suas capacidades analíticas, a organização procura "
        "um(a) Analista de Dados Sénior para liderar a análise de dados programáticos, "
        "desenvolver dashboards de monitoria e apoiar a tomada de decisão baseada em evidências."
    )

    # Objectives
    set_heading(doc, "2. Objectivos da Posição", level=2, color=(31, 73, 125))
    for obj in [
        "Liderar a análise e visualização de dados programáticos de saúde a nível nacional e provincial.",
        "Desenvolver e manter dashboards de monitoria em Power BI e DHIS2.",
        "Garantir a qualidade e integridade dos dados reportados pelas unidades sanitárias.",
        "Apoiar a equipa de monitoria e avaliação no desenvolvimento de relatórios periódicos.",
        "Construir capacidades analíticas das equipas provinciais através de formações regulares.",
    ]:
        doc.add_paragraph(f"• {obj}", style="List Bullet")

    # Responsibilities
    set_heading(doc, "3. Responsabilidades Principais", level=2, color=(31, 73, 125))
    responsabilidades = [
        ("Análise de Dados (40%)",
         ["Processar e analisar dados mensais do SISMA e DHIS2 para os 6 indicadores prioritários do PEPFAR.",
          "Realizar análises estatísticas de tendências, outliers e qualidade de dados.",
          "Produzir relatórios analíticos mensais, trimestrais e anuais."]),
        ("Visualização e Reporting (30%)",
         ["Desenvolver e actualizar dashboards interactivos em Power BI.",
          "Criar apresentações executivas para doadores e parceiros do Governo.",
          "Apoiar a preparação do relatório anual PEPFAR/USAID (APR)."]),
        ("Gestão de Dados (20%)",
         ["Garantir a limpeza, validação e integridade dos dados recebidos.",
          "Gerir a base de dados central da organização (PostgreSQL/Excel).",
          "Documentar processos e procedimentos de gestão de dados."]),
        ("Capacitação (10%)",
         ["Facilitar formações em análise de dados para equipas provinciais.",
          "Desenvolver guias e materiais de treinamento."]),
    ]
    for titulo_resp, items in responsabilidades:
        p = doc.add_paragraph()
        p.add_run(titulo_resp).bold = True
        for item in items:
            doc.add_paragraph(f"• {item}", style="List Bullet")
        doc.add_paragraph()

    # Requirements
    set_heading(doc, "4. Requisitos da Posição", level=2, color=(31, 73, 125))

    set_heading(doc, "4.1 Formação Académica", level=3)
    doc.add_paragraph(
        "Mestrado em Saúde Pública, Epidemiologia, Estatística, Ciência de Dados ou área relacionada. "
        "Licenciatura com experiência adicional relevante poderá ser considerada."
    )

    set_heading(doc, "4.2 Experiência", level=3)
    for exp in [
        "Mínimo de 5 anos de experiência em análise de dados, preferencialmente em saúde pública.",
        "Experiência comprovada com DHIS2, SISMA ou sistemas similares de informação de saúde.",
        "Experiência de trabalho com programas financiados por PEPFAR ou USAID (preferencial).",
    ]:
        doc.add_paragraph(f"• {exp}", style="List Bullet")

    set_heading(doc, "4.3 Competências Técnicas Obrigatórias", level=3)
    for comp in [
        "Python (pandas, numpy, matplotlib) — nível avançado",
        "SQL — nível avançado (PostgreSQL ou MySQL)",
        "Power BI — nível avançado",
        "Excel — nível avançado",
        "Machine Learning — nível intermédio",
        "DHIS2 — nível intermédio",
        "Gestão de Projetos",
    ]:
        doc.add_paragraph(f"• {comp}", style="List Bullet")

    set_heading(doc, "4.4 Competências Transversais", level=3)
    for soft in [
        "Comunicação eficaz oral e escrita em Português e Inglês",
        "Capacidade de trabalhar em ambiente multicultural e de alta pressão",
        "Pensamento analítico e resolução de problemas",
        "Liderança e trabalho em equipa",
    ]:
        doc.add_paragraph(f"• {soft}", style="List Bullet")

    # Process
    set_heading(doc, "5. Processo de Candidatura", level=2, color=(31, 73, 125))
    doc.add_paragraph(
        "Os candidatos interessados devem submeter os seguintes documentos até 31 de Julho de 2026:"
    )
    for doc_item in [
        "Carta de motivação (máximo 1 página)",
        "Curriculum Vitae actualizado (máximo 4 páginas)",
        "Cópias dos certificados académicos",
        "Contactos de 3 referências profissionais",
    ]:
        doc.add_paragraph(f"• {doc_item}", style="List Bullet")

    doc.add_paragraph()
    doc.add_paragraph(
        "Candidaturas devem ser enviadas para: recrutamento@talentiq.co.mz com o assunto "
        "\"Analista de Dados Sénior — [Seu Nome]\"."
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(
        "Apenas candidatos seleccionados para entrevista serão contactados. "
        "TalentIQ Health Partners é um empregador de igualdade de oportunidades."
    ).italic = True

    path = os.path.join(OUTPUT_DIR, "ToR_Analista_Dados_Senior.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# ToR — Especialista de RH
# ─────────────────────────────────────────────
def criar_tor_especialista_rh():
    doc = Document()

    title = doc.add_heading("TERMOS DE REFERÊNCIA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("Especialista de Recursos Humanos", 1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    cells = [
        ("Organização", "TalentIQ Health Partners"),
        ("Localização", "Maputo, Moçambique"),
        ("Duração do Contrato", "24 meses renováveis"),
        ("Data Limite de Candidatura", "15 de Agosto de 2026"),
    ]
    for i, (label, value) in enumerate(cells):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    set_heading(doc, "1. Contexto", level=2, color=(31, 73, 125))
    doc.add_paragraph(
        "A TalentIQ Health Partners procura um(a) Especialista de Recursos Humanos para "
        "liderar os processos de recrutamento, gestão de desempenho e desenvolvimento "
        "organizacional da organização, que conta actualmente com 85 colaboradores directos "
        "e mais de 200 consultores parceiros em 5 províncias."
    )

    set_heading(doc, "2. Responsabilidades Principais", level=2, color=(31, 73, 125))
    for r in [
        "Gerir o ciclo completo de recrutamento e seleção para todas as posições da organização.",
        "Desenvolver e implementar políticas e procedimentos de RH alinhados com a legislação laboral.",
        "Coordenar o processo de avaliação de desempenho semestral e anual.",
        "Gerir programas de onboarding, integração e desenvolvimento de colaboradores.",
        "Produzir relatórios de RH para a liderança e doadores.",
        "Garantir conformidade com legislação laboral moçambicana e políticas de doadores.",
    ]:
        doc.add_paragraph(f"• {r}", style="List Bullet")

    set_heading(doc, "3. Requisitos", level=2, color=(31, 73, 125))

    set_heading(doc, "3.1 Formação", level=3)
    doc.add_paragraph(
        "MBA ou Mestrado em Gestão de Recursos Humanos, Gestão de Empresas ou área relacionada. "
        "Licenciatura com experiência mínima de 8 anos será considerada."
    )

    set_heading(doc, "3.2 Experiência e Competências", level=3)
    for c in [
        "Mínimo de 5 anos em posições de RH, com pelo menos 2 em nível sénior.",
        "Experiência em recrutamento e seleção de posições técnicas e de gestão.",
        "Domínio de Excel, Power BI para relatórios de RH.",
        "Conhecimento de ferramentas de gestão de RH (SAP HR, BambooHR ou similar).",
        "Competências em negociação, liderança e comunicação.",
        "Fluência em Português e Inglês (obrigatório).",
        "Experiência em organizações internacionais de saúde ou desenvolvimento (preferencial).",
    ]:
        doc.add_paragraph(f"• {c}", style="List Bullet")

    set_heading(doc, "4. Processo de Candidatura", level=2, color=(31, 73, 125))
    doc.add_paragraph("Enviar CV e carta de motivação para: recrutamento@talentiq.co.mz")
    doc.add_paragraph("Prazo: 15 de Agosto de 2026")

    path = os.path.join(OUTPUT_DIR, "ToR_Especialista_RH.docx")
    doc.save(path)
    print(f"✅ {path}")


# ─────────────────────────────────────────────
# Run all
# ─────────────────────────────────────────────
if __name__ == "__main__":
    print("\n🚀 A gerar documentos de teste para TalentIQ...\n")
    criar_cv_ana_macie()
    criar_cv_carlos_sitoe()
    criar_cv_fatima_cumbe()
    criar_cv_joao_nhantumbo()
    criar_tor_analista_dados()
    criar_tor_especialista_rh()
    print(f"\n✅ Todos os documentos gerados em: {OUTPUT_DIR}\n")
    print("Documentos disponíveis:")
    for f in sorted(os.listdir(OUTPUT_DIR)):
        print(f"  📄 {f}")
